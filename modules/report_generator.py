# modules/report_generator.py
from typing import List, Dict, Optional, Union, Tuple
from datetime import datetime
from dataclasses import dataclass, asdict
from enum import Enum
import json
import io
import base64
from pathlib import Path
from modules.meeting_context import MeetingContext

# Core libraries
from modules.logger import logger
import arrow
import streamlit as st

# Visualization libraries
try:
    import matplotlib.pyplot as plt
    import matplotlib.patches as patches
    import seaborn as sns
    MATPLOTLIB_AVAILABLE = True
except ImportError:
    logger.warning("Matplotlib not available - visual reports disabled")
    MATPLOTLIB_AVAILABLE = False

try:
    import plotly.graph_objects as go
    import plotly.express as px
    from plotly.subplots import make_subplots
    PLOTLY_AVAILABLE = True
except ImportError:
    logger.warning("Plotly not available - interactive charts disabled")
    PLOTLY_AVAILABLE = False

# Export libraries
try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
    REPORTLAB_AVAILABLE = True
except ImportError:
    logger.warning("ReportLab not available - PDF export disabled")
    REPORTLAB_AVAILABLE = False

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available - Word export disabled")
    DOCX_AVAILABLE = False

try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    logger.warning("Pandas not available - Excel export disabled")
    PANDAS_AVAILABLE = False

class ExportFormat(Enum):
    TXT = "txt"
    PDF = "pdf"
    DOCX = "docx"
    HTML = "html"
    JSON = "json"
    EXCEL = "xlsx"

class VisualizationType(Enum):
    SPEAKER_CONTRIBUTION = "speaker_contribution"
    ACTION_PRIORITY = "action_priority"
    DEADLINE_TIMELINE = "deadline_timeline"
    WORD_CLOUD = "word_cloud"
    SENTIMENT_ANALYSIS = "sentiment_analysis"

@dataclass
class ReportConfig:
    """Configuration for report generation"""
    include_visuals: bool = True
    include_transcript: bool = True
    include_statistics: bool = True
    include_speaker_analysis: bool = True
    include_priority_analysis: bool = True
    include_timeline: bool = True
    visual_style: str = "professional"  # professional, modern, colorful
    page_orientation: str = "portrait"  # portrait, landscape
    font_size: int = 11
    export_format: ExportFormat = ExportFormat.TXT
    
    def to_dict(self) -> Dict:
        return {**asdict(self), 'export_format': self.export_format.value}

@dataclass 
class ReportData:
    """Enhanced data structure for report generation"""
    participants: List[Dict] = None  # From enhanced NER
    action_items: List[Dict] = None  # From enhanced NER  
    dates: List[Dict] = None         # From enhanced NER
    summary: str = ""
    transcript: str = ""
    speaker_segments: List[Tuple[str, str]] = None
    title: str = "Meeting Summary Report"
    metadata: Dict = None
    
    def __post_init__(self):
        if self.participants is None:
            self.participants = []
        if self.action_items is None:
            self.action_items = []
        if self.dates is None:
            self.dates = []
        if self.speaker_segments is None:
            self.speaker_segments = []
        if self.metadata is None:
            self.metadata = {}

class EnhancedReportGenerator:
    """Enhanced report generator with visuals and multiple export formats"""
    
    def __init__(self):
        self.color_palette = {
            'professional': ['#2E4A62', '#8D9DB6', '#F8F8F8', '#4A4A4A'],
            'modern': ['#FF6B6B', '#4ECDC4', '#45B7D1', '#96CEB4'],
            'colorful': ['#FF9F43', '#10AC84', '#EE5A24', '#0ABDE3']
        }
    
    def _analyze_speaker_contributions(self, speaker_segments: List[Tuple[str, str]]) -> Dict:
        """Analyze speaker contributions and statistics"""
        if not speaker_segments:
            return {}
        
        speaker_stats = {}
        total_words = 0
        
        for speaker, text in speaker_segments:
            words = len(text.split())
            total_words += words
            
            if speaker in speaker_stats:
                speaker_stats[speaker]['word_count'] += words
                speaker_stats[speaker]['segment_count'] += 1
            else:
                speaker_stats[speaker] = {
                    'word_count': words,
                    'segment_count': 1,
                    'avg_words_per_segment': 0,
                    'contribution_percentage': 0
                }
        
        # Calculate percentages and averages
        for speaker in speaker_stats:
            stats = speaker_stats[speaker]
            stats['contribution_percentage'] = (stats['word_count'] / max(total_words, 1)) * 100
            stats['avg_words_per_segment'] = stats['word_count'] / max(stats['segment_count'], 1)
        
        return {
            'speaker_stats': speaker_stats,
            'total_speakers': len(speaker_stats),
            'total_words': total_words,
            'most_active': max(speaker_stats.keys(), key=lambda x: speaker_stats[x]['word_count']) if speaker_stats else None
        }
    
    def _analyze_action_priorities(self, action_items: List[Dict]) -> Dict:
        """Analyze action item priorities and assignments"""
        if not action_items:
            return {}
        
        priority_counts = {'high': 0, 'medium': 0, 'low': 0, 'unspecified': 0}
        assignee_counts = {}
        methods_used = {}
        
        for action in action_items:
            # Count priorities
            priority = action.get('priority', 'unspecified')
            priority_counts[priority] = priority_counts.get(priority, 0) + 1
            
            # Count assignees
            assignee = action.get('assignee', 'Unassigned')
            assignee_counts[assignee] = assignee_counts.get(assignee, 0) + 1
            
            # Count extraction methods
            method = action.get('extraction_method', 'unknown')
            methods_used[method] = methods_used.get(method, 0) + 1
        
        return {
            'priority_distribution': priority_counts,
            'assignee_distribution': assignee_counts,
            'extraction_methods': methods_used,
            'total_actions': len(action_items),
            'assigned_actions': len([a for a in action_items if a.get('assignee') and a.get('assignee') != 'Unassigned']),
            'high_priority_count': priority_counts.get('high', 0)
        }
    
    def _create_speaker_contribution_chart(self, speaker_analysis: Dict, style: str = "professional") -> Optional[str]:
        """Create speaker contribution pie chart"""
        if not MATPLOTLIB_AVAILABLE or not speaker_analysis.get('speaker_stats'):
            return None
        
        try:
            fig, ax = plt.subplots(figsize=(10, 8))
            
            speakers = list(speaker_analysis['speaker_stats'].keys())
            contributions = [speaker_analysis['speaker_stats'][s]['contribution_percentage'] for s in speakers]
            colors = self.color_palette[style][:len(speakers)]
            
            # Create pie chart
            wedges, texts, autotexts = ax.pie(
                contributions, 
                labels=speakers, 
                autopct='%1.1f%%',
                colors=colors,
                startangle=90,
                textprops={'fontsize': 12}
            )
            
            ax.set_title('Speaker Contribution Distribution', fontsize=16, fontweight='bold', pad=20)
            
            # Add statistics box
            stats_text = f"""Total Speakers: {speaker_analysis['total_speakers']}
Total Words: {speaker_analysis['total_words']:,}
Most Active: {speaker_analysis.get('most_active', 'N/A')}"""
            
            ax.text(1.3, 0.5, stats_text, transform=ax.transAxes, fontsize=10,
                   bbox=dict(boxstyle="round,pad=0.3", facecolor='lightgray', alpha=0.5))
            
            plt.tight_layout()
            
            # Convert to base64 string
            buffer = io.BytesIO()
            plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
            buffer.seek(0)
            chart_base64 = base64.b64encode(buffer.getvalue()).decode()
            plt.close()
            
            return chart_base64
            
        except Exception as e:
            logger.exception("Failed to create speaker contribution chart")
            return None
    
    def _create_priority_distribution_chart(self, priority_analysis: Dict, style: str = "professional") -> Optional[str]:
        """Create action item priority distribution chart"""
        if not MATPLOTLIB_AVAILABLE or not priority_analysis.get('priority_distribution'):
            return None
        
        try:
            fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(15, 6))
            
            priorities = list(priority_analysis['priority_distribution'].keys())
            counts = list(priority_analysis['priority_distribution'].values())
            colors = self.color_palette[style][:len(priorities)]
            
            # Priority distribution bar chart
            bars = ax1.bar(priorities, counts, color=colors)
            ax1.set_title('Action Items by Priority', fontsize=14, fontweight='bold')
            ax1.set_ylabel('Number of Items')
            ax1.set_xlabel('Priority Level')
            
            # Add value labels on bars
            for bar in bars:
                height = bar.get_height()
                ax1.text(bar.get_x() + bar.get_width()/2., height + 0.1,
                        f'{int(height)}', ha='center', va='bottom')
            
            # Assignee distribution pie chart
            if priority_analysis.get('assignee_distribution'):
                assignees = list(priority_analysis['assignee_distribution'].keys())
                assignee_counts = list(priority_analysis['assignee_distribution'].values())
                
                ax2.pie(assignee_counts, labels=assignees, autopct='%1.1f%%', 
                       colors=colors, startangle=90)
                ax2.set_title('Action Items by Assignee', fontsize=14, fontweight='bold')
            
            plt.tight_layout()
            
            # Convert to base64 string
            buffer = io.BytesIO()
            plt.savefig(buffer, format='png', dpi=300, bbox_inches='tight')
            buffer.seek(0)
            chart_base64 = base64.b64encode(buffer.getvalue()).decode()
            plt.close()
            
            return chart_base64
            
        except Exception as e:
            logger.exception("Failed to create priority distribution chart")
            return None
    
    def _create_plotly_dashboard(self, report_data: ReportData, config: ReportConfig) -> Optional[str]:
        """Create interactive Plotly dashboard"""
        if not PLOTLY_AVAILABLE:
            return None
        
        try:
            # Analyze data
            speaker_analysis = self._analyze_speaker_contributions(report_data.speaker_segments)
            priority_analysis = self._analyze_action_priorities(report_data.action_items)
            
            # Create subplots
            fig = make_subplots(
                rows=2, cols=2,
                specs=[[{"type": "pie"}, {"type": "bar"}],
                       [{"type": "table"}, {"type": "scatter"}]],
                subplot_titles=('Speaker Contributions', 'Action Priority Distribution',
                              'Action Items Summary', 'Meeting Timeline')
            )
            
            # Speaker contribution pie chart
            if speaker_analysis.get('speaker_stats'):
                speakers = list(speaker_analysis['speaker_stats'].keys())
                contributions = [speaker_analysis['speaker_stats'][s]['contribution_percentage'] for s in speakers]
                
                fig.add_trace(go.Pie(
                    labels=speakers, 
                    values=contributions,
                    name="Speaker Contributions"
                ), row=1, col=1)
            
            # Priority distribution bar chart
            if priority_analysis.get('priority_distribution'):
                priorities = list(priority_analysis['priority_distribution'].keys())
                counts = list(priority_analysis['priority_distribution'].values())
                
                fig.add_trace(go.Bar(
                    x=priorities,
                    y=counts,
                    name="Priority Distribution"
                ), row=1, col=2)
            
            # Action items table
            if report_data.action_items:
                table_data = []
                for action in report_data.action_items[:10]:  # Show first 10
                    table_data.append([
                        action.get('text', '')[:50] + '...' if len(action.get('text', '')) > 50 else action.get('text', ''),
                        action.get('priority', 'N/A'),
                        action.get('assignee', 'N/A'),
                        action.get('deadline', 'N/A')
                    ])
                
                fig.add_trace(go.Table(
                    header=dict(values=['Action Item', 'Priority', 'Assignee', 'Deadline']),
                    cells=dict(values=list(map(list, zip(*table_data))))
                ), row=2, col=1)
            
            fig.update_layout(
                title_text="Meeting Analysis Dashboard",
                showlegend=True,
                height=800
            )
            
            # Convert to HTML
            html_content = fig.to_html(include_plotlyjs='cdn', div_id="meeting-dashboard")
            return html_content
            
        except Exception as e:
            logger.exception("Failed to create Plotly dashboard")
            return None
    
    def _generate_text_report(self, report_data: ReportData, config: ReportConfig,
                            speaker_analysis: Dict, priority_analysis: Dict) -> str:
        """Generate enhanced text report"""
        try:
            report_lines = []
            
            # Title and header
            title = report_data.title or "Meeting Summary Report"
            report_lines.append(title)
            report_lines.append("=" * len(title))
            
            timestamp = arrow.now().format("YYYY-MM-DD HH:mm A (ZZ)")
            report_lines.append(f"Generated on: {timestamp}")
            report_lines.append("")
            
            # Executive Summary
            if config.include_statistics:
                report_lines.append("📊 Executive Summary")
                report_lines.append("-" * 50)
                
                if speaker_analysis:
                    report_lines.append(f"• Total Participants: {speaker_analysis.get('total_speakers', 0)}")
                    report_lines.append(f"• Most Active Speaker: {speaker_analysis.get('most_active', 'N/A')}")
                
                if priority_analysis:
                    report_lines.append(f"• Total Action Items: {priority_analysis.get('total_actions', 0)}")
                    report_lines.append(f"• High Priority Actions: {priority_analysis.get('high_priority_count', 0)}")
                    report_lines.append(f"• Assigned Actions: {priority_analysis.get('assigned_actions', 0)}")
                
                report_lines.append("")
            
            # Metadata
            if report_data.metadata:
                report_lines.append("📋 Meeting Details")
                report_lines.append("-" * 50)
                for key, value in report_data.metadata.items():
                    report_lines.append(f"• {key}: {value}")
                report_lines.append("")
            
            # Participants with roles
            report_lines.append("👥 Participants")
            report_lines.append("-" * 50)
            if report_data.participants:
                for participant in report_data.participants:
                    if isinstance(participant, dict):
                        name = participant.get('name', 'Unknown')
                        roles = participant.get('roles', [])
                        role_text = f" ({', '.join(roles)})" if roles else ""
                        mentions = participant.get('mentions', 0)
                        report_lines.append(f"• {name}{role_text} - {mentions} mentions")
                    else:
                        report_lines.append(f"• {participant}")
            else:
                report_lines.append("No participants detected.")
            report_lines.append("")
            
            # Speaker Analysis
            if config.include_speaker_analysis and speaker_analysis.get('speaker_stats'):
                report_lines.append("🎤 Speaker Analysis")
                report_lines.append("-" * 50)
                for speaker, stats in speaker_analysis['speaker_stats'].items():
                    report_lines.append(f"• {speaker}:")
                    report_lines.append(f"  - Contribution: {stats['contribution_percentage']:.1f}%")
                    report_lines.append(f"  - Word Count: {stats['word_count']:,}")
                    report_lines.append(f"  - Segments: {stats['segment_count']}")
                    report_lines.append(f"  - Avg Words/Segment: {stats['avg_words_per_segment']:.1f}")
                report_lines.append("")
            
            # Action Items with Priority
            report_lines.append("📌 Action Items")
            report_lines.append("-" * 50)
            if report_data.action_items:
                # Group by priority
                priority_groups = {'high': [], 'medium': [], 'low': [], 'unspecified': []}
                for action in report_data.action_items:
                    if isinstance(action, dict):
                        priority = action.get('priority', 'unspecified')
                        priority_groups[priority].append(action)
                    else:
                        priority_groups['unspecified'].append({'text': action})
                
                for priority in ['high', 'medium', 'low', 'unspecified']:
                    if priority_groups[priority]:
                        priority_symbol = {'high': '🔴', 'medium': '🟡', 'low': '🟢', 'unspecified': '⚪'}
                        report_lines.append(f"\n{priority_symbol[priority]} {priority.upper()} PRIORITY:")
                        
                        for action in priority_groups[priority]:
                            text = action.get('text', action) if isinstance(action, dict) else action
                            assignee = action.get('assignee', '') if isinstance(action, dict) else ''
                            deadline = action.get('deadline', '') if isinstance(action, dict) else ''
                            
                            action_line = f"• {text}"
                            if assignee and assignee != 'Unassigned':
                                action_line += f" [Assignee: {assignee}]"
                            if deadline:
                                action_line += f" [Due: {deadline}]"
                            
                            report_lines.append(action_line)
            else:
                report_lines.append("No action items found.")
            report_lines.append("")
            
            # Priority Analysis
            if config.include_priority_analysis and priority_analysis:
                report_lines.append("📈 Priority Analysis")
                report_lines.append("-" * 50)
                dist = priority_analysis.get('priority_distribution', {})
                total = sum(dist.values())
                
                for priority, count in dist.items():
                    percentage = (count / max(total, 1)) * 100
                    report_lines.append(f"• {priority.title()}: {count} items ({percentage:.1f}%)")
                
                if priority_analysis.get('assignee_distribution'):
                    report_lines.append("\nAssignee Distribution:")
                    for assignee, count in priority_analysis['assignee_distribution'].items():
                        report_lines.append(f"• {assignee}: {count} items")
                report_lines.append("")
            
            # Deadlines and Dates
            report_lines.append("⏰ Deadlines & Important Dates")
            report_lines.append("-" * 50)
            if report_data.dates:
                for date_item in report_data.dates:
                    if isinstance(date_item, dict):
                        text = date_item.get('text', '')
                        context = date_item.get('context', '')
                        if context:
                            report_lines.append(f"• {text} - {context}")
                        else:
                            report_lines.append(f"• {text}")
                    else:
                        report_lines.append(f"• {date_item}")
            else:
                report_lines.append("No dates or deadlines detected.")
            report_lines.append("")
            
            # Meeting Summary
            report_lines.append("🧠 Meeting Summary")
            report_lines.append("-" * 50)
            report_lines.append(report_data.summary.strip() or "No summary available.")
            report_lines.append("")
            
            # Full Transcript
            if config.include_transcript and report_data.transcript:
                report_lines.append("📄 Full Transcript")
                report_lines.append("-" * 50)
                report_lines.append(report_data.transcript.strip())
            
            return "\n".join(report_lines)
            
        except Exception as e:
            logger.exception("Failed to generate text report")
            return "Error generating text report"
    
    def _export_to_pdf(self, content: str, title: str) -> Optional[bytes]:
        """Export report to PDF format"""
        if not REPORTLAB_AVAILABLE:
            return None
        
        try:
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter, topMargin=0.75*inch)
            
            # Styles
            styles = getSampleStyleSheet()
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                alignment=TA_CENTER,
                textColor=colors.HexColor('#2E4A62')
            )
            
            heading_style = ParagraphStyle(
                'CustomHeading',
                parent=styles['Heading2'],
                fontSize=14,
                textColor=colors.HexColor('#4A4A4A')
            )
            
            # Content
            story = []
            lines = content.split('\n')
            
            for line in lines:
                if line.strip():
                    if line.startswith('='):
                        continue  # Skip separator lines
                    elif line.endswith('Report') or 'Summary' in line:
                        story.append(Paragraph(line, title_style))
                    elif line.startswith('Generated on:'):
                        story.append(Paragraph(line, styles['Normal']))
                    elif any(line.startswith(prefix) for prefix in ['📊', '👥', '🎤', '📌', '📈', '⏰', '🧠', '📄']):
                        story.append(Spacer(1, 12))
                        story.append(Paragraph(line, heading_style))
                    elif line.startswith('-'):
                        continue  # Skip separator lines
                    else:
                        story.append(Paragraph(line, styles['Normal']))
                else:
                    story.append(Spacer(1, 6))
            
            doc.build(story)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.exception("Failed to export to PDF")
            return None
    
    def _export_to_docx(self, content: str, title: str) -> Optional[bytes]:
        """Export report to Word document"""
        if not DOCX_AVAILABLE:
            return None
        
        try:
            doc = Document()
            
            # Add title
            title_para = doc.add_heading(title, 0)
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Process content
            lines = content.split('\n')
            for line in lines:
                if line.strip():
                    if line.startswith('=') or line.startswith('-'):
                        continue  # Skip separator lines
                    elif any(line.startswith(prefix) for prefix in ['📊', '👥', '🎤', '📌', '📈', '⏰', '🧠', '📄']):
                        doc.add_heading(line, level=1)
                    elif line.startswith('•'):
                        doc.add_paragraph(line, style='List Bullet')
                    else:
                        doc.add_paragraph(line)
                else:
                    doc.add_paragraph()  # Empty line
            
            # Save to buffer
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.exception("Failed to export to DOCX")
            return None
    
    def _export_to_excel(self, report_data: ReportData, speaker_analysis: Dict, priority_analysis: Dict) -> Optional[bytes]:
        """Export structured data to Excel"""
        if not PANDAS_AVAILABLE:
            return None
        
        try:
            buffer = io.BytesIO()
            
            with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
                # Summary sheet
                summary_data = {
                    'Metric': ['Total Participants', 'Total Action Items', 'High Priority Actions', 'Meeting Date'],
                    'Value': [
                        speaker_analysis.get('total_speakers', 0),
                        priority_analysis.get('total_actions', 0),
                        priority_analysis.get('high_priority_count', 0),
                        arrow.now().format('YYYY-MM-DD')
                    ]
                }
                pd.DataFrame(summary_data).to_excel(writer, sheet_name='Summary', index=False)
                
                # Action Items sheet
                if report_data.action_items:
                    actions_df = pd.DataFrame([
                        {
                            'Action Item': action.get('text', '') if isinstance(action, dict) else str(action),
                            'Priority': action.get('priority', '') if isinstance(action, dict) else '',
                            'Assignee': action.get('assignee', '') if isinstance(action, dict) else '',
                            'Deadline': action.get('deadline', '') if isinstance(action, dict) else '',
                            'Confidence': action.get('confidence', '') if isinstance(action, dict) else ''
                        } for action in report_data.action_items
                    ])
                    actions_df.to_excel(writer, sheet_name='Action Items', index=False)
                
                # Participants sheet
                if report_data.participants:
                    participants_df = pd.DataFrame([
                        {
                            'Name': p.get('name', '') if isinstance(p, dict) else str(p),
                            'Mentions': p.get('mentions', '') if isinstance(p, dict) else '',
                            'Roles': ', '.join(p.get('roles', [])) if isinstance(p, dict) else ''
                        } for p in report_data.participants
                    ])
                    participants_df.to_excel(writer, sheet_name='Participants', index=False)
                
                # Speaker Analysis sheet
                if speaker_analysis.get('speaker_stats'):
                    speaker_df = pd.DataFrame([
                        {
                            'Speaker': speaker,
                            'Word Count': stats['word_count'],
                            'Contribution %': round(stats['contribution_percentage'], 1),
                            'Segments': stats['segment_count'],
                            'Avg Words/Segment': round(stats['avg_words_per_segment'], 1)
                        } for speaker, stats in speaker_analysis['speaker_stats'].items()
                    ])
                    speaker_df.to_excel(writer, sheet_name='Speaker Analysis', index=False)
            
            buffer.seek(0)
            return buffer.getvalue()
            
        except Exception as e:
            logger.exception("Failed to export to Excel")
            return None
    
    def generate_report(self, report_data: ReportData, config: ReportConfig = None) -> Dict:
        """
        Generate comprehensive meeting report with analytics and visualizations
        """
        if config is None:
            config = ReportConfig()
        
        try:
            logger.info(f"🚀 Generating {config.export_format.value.upper()} report")
            
            # Analyze data
            speaker_analysis = self._analyze_speaker_contributions(report_data.speaker_segments)
            priority_analysis = self._analyze_action_priorities(report_data.action_items)
            
            result = {
                'success': True,
                'format': config.export_format.value,
                'content': None,
                'filename': None,
                'charts': {},
                'analytics': {
                    'speaker_analysis': speaker_analysis,
                    'priority_analysis': priority_analysis
                },
                'timestamp': arrow.now().format('YYYY-MM-DD_HH-mm')
            }
            
            # Generate visualizations
            if config.include_visuals and MATPLOTLIB_AVAILABLE:
                result['charts']['speaker_contribution'] = self._create_speaker_contribution_chart(
                    speaker_analysis, config.visual_style
                )
                result['charts']['priority_distribution'] = self._create_priority_distribution_chart(
                    priority_analysis, config.visual_style
                )
            
            # Generate interactive dashboard
            if config.include_visuals and PLOTLY_AVAILABLE:
                result['charts']['interactive_dashboard'] = self._create_plotly_dashboard(report_data, config)
            
            # Generate content based on format
            title = report_data.title or "Meeting Summary Report"
            timestamp = result['timestamp']
            
            if config.export_format == ExportFormat.TXT:
                result['content'] = self._generate_text_report(report_data, config, speaker_analysis, priority_analysis)
                result['filename'] = f"meeting_report_{timestamp}.txt"
                result['mime_type'] = "text/plain"
                
            elif config.export_format == ExportFormat.PDF:
                text_content = self._generate_text_report(report_data, config, speaker_analysis, priority_analysis)
                result['content'] = self._export_to_pdf(text_content, title)
                result['filename'] = f"meeting_report_{timestamp}.pdf"
                result['mime_type'] = "application/pdf"
                
            elif config.export_format == ExportFormat.DOCX:
                text_content = self._generate_text_report(report_data, config, speaker_analysis, priority_analysis)
                result['content'] = self._export_to_docx(text_content, title)
                result['filename'] = f"meeting_report_{timestamp}.docx"
                result['mime_type'] = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                
            elif config.export_format == ExportFormat.HTML:
                text_content = self._generate_text_report(report_data, config, speaker_analysis, priority_analysis)
                html_content = self._generate_html_report(text_content, result['charts'], config)
                result['content'] = html_content
                result['filename'] = f"meeting_report_{timestamp}.html"
                result['mime_type'] = "text/html"
                
            elif config.export_format == ExportFormat.JSON:
                json_content = self._generate_json_report(report_data, speaker_analysis, priority_analysis)
                result['content'] = json_content
                result['filename'] = f"meeting_report_{timestamp}.json"
                result['mime_type'] = "application/json"
                
            elif config.export_format == ExportFormat.EXCEL:
                result['content'] = self._export_to_excel(report_data, speaker_analysis, priority_analysis)
                result['filename'] = f"meeting_report_{timestamp}.xlsx"
                result['mime_type'] = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
            
            if result['content'] is None:
                result['success'] = False
                result['error'] = f"Failed to generate {config.export_format.value} content"
            
            logger.info(f"✅ Report generated successfully: {result['filename']}")
            return result
            
        except Exception as e:
            logger.exception("❌ Failed to generate report")
            return {
                'success': False,
                'error': str(e),
                'format': config.export_format.value if config else 'unknown'
            }
    
    def _generate_html_report(self, text_content: str, charts: Dict, config: ReportConfig) -> str:
        """Generate HTML report with embedded visualizations"""
        try:
            html_template = """
<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Meeting Summary Report</title>
    <style>
        body {{ font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif; margin: 40px; background-color: #f5f5f5; }}
        .container {{ max-width: 1200px; margin: 0 auto; background: white; padding: 30px; border-radius: 10px; box-shadow: 0 0 20px rgba(0,0,0,0.1); }}
        h1 {{ color: #2E4A62; text-align: center; border-bottom: 3px solid #2E4A62; padding-bottom: 10px; }}
        h2 {{ color: #4A4A4A; border-left: 4px solid #8D9DB6; padding-left: 15px; margin-top: 30px; }}
        .chart-container {{ text-align: center; margin: 20px 0; padding: 20px; background: #f9f9f9; border-radius: 8px; }}
        .chart-container img {{ max-width: 100%; height: auto; border-radius: 8px; box-shadow: 0 4px 8px rgba(0,0,0,0.1); }}
        .stats-box {{ background: linear-gradient(135deg, #667eea 0%, #764ba2 100%); color: white; padding: 20px; border-radius: 10px; margin: 20px 0; }}
        .priority-high {{ color: #e74c3c; font-weight: bold; }}
        .priority-medium {{ color: #f39c12; font-weight: bold; }}
        .priority-low {{ color: #27ae60; font-weight: bold; }}
        ul {{ list-style-type: none; padding-left: 0; }}
        li {{ margin: 8px 0; padding: 8px; background: #f8f9fa; border-left: 3px solid #dee2e6; border-radius: 4px; }}
        .metadata {{ background: #e8f4f8; padding: 15px; border-radius: 8px; margin: 20px 0; }}
        pre {{ background: #f4f4f4; padding: 15px; border-radius: 8px; overflow-x: auto; }}
    </style>
</head>
<body>
    <div class="container">
        {content}
        
        {charts_section}
        
        {interactive_dashboard}
    </div>
</body>
</html>
"""
            
            # Convert text content to HTML
            html_content = text_content.replace('\n', '<br>')
            
            # Add chart sections
            charts_html = ""
            if charts.get('speaker_contribution'):
                charts_html += f'''
                <div class="chart-container">
                    <h3>Speaker Contribution Analysis</h3>
                    <img src="data:image/png;base64,{charts['speaker_contribution']}" alt="Speaker Contribution Chart">
                </div>
                '''
            
            if charts.get('priority_distribution'):
                charts_html += f'''
                <div class="chart-container">
                    <h3>Action Items Priority Distribution</h3>
                    <img src="data:image/png;base64,{charts['priority_distribution']}" alt="Priority Distribution Chart">
                </div>
                '''
            
            # Add interactive dashboard
            dashboard_html = ""
            if charts.get('interactive_dashboard'):
                dashboard_html = f'''
                <div class="chart-container">
                    <h2>Interactive Dashboard</h2>
                    {charts['interactive_dashboard']}
                </div>
                '''
            
            return html_template.format(
                content=html_content,
                charts_section=charts_html,
                interactive_dashboard=dashboard_html
            )
            
        except Exception as e:
            logger.exception("Failed to generate HTML report")
            return f"<html><body><h1>Error generating HTML report</h1><p>{str(e)}</p></body></html>"
    
    def _generate_json_report(self, report_data: ReportData, speaker_analysis: Dict, priority_analysis: Dict) -> str:
        """Generate structured JSON report"""
        try:
            json_data = {
                'metadata': {
                    'title': report_data.title,
                    'generated_at': arrow.now().isoformat(),
                    'generator_version': '2.0.0'
                },
                'summary': report_data.summary,
                'participants': report_data.participants,
                'action_items': report_data.action_items,
                'dates_and_deadlines': report_data.dates,
                'analytics': {
                    'speaker_analysis': speaker_analysis,
                    'priority_analysis': priority_analysis
                },
                'speaker_segments': [
                    {'speaker': speaker, 'text': text, 'word_count': len(text.split())}
                    for speaker, text in report_data.speaker_segments
                ],
                'transcript': report_data.transcript if len(report_data.transcript) < 50000 else report_data.transcript[:50000] + "...[truncated]"
            }
            
            return json.dumps(json_data, indent=2, ensure_ascii=False)
            
        except Exception as e:
            logger.exception("Failed to generate JSON report")
            return json.dumps({'error': f'Failed to generate JSON report: {str(e)}'}, indent=2)

# Global report generator instance
_report_generator_instance = None

def get_report_generator() -> EnhancedReportGenerator:
    """Get singleton report generator instance"""
    global _report_generator_instance
    if _report_generator_instance is None:
        _report_generator_instance = EnhancedReportGenerator()
    return _report_generator_instance

# Backward compatible function
def generate_meeting_report(participants: List[str],
                          action_items: List[str], 
                          deadlines: List[str],
                          summary: str,
                          original_transcript: str,
                          title: str = "Meeting Summary Report",
                          metadata: Optional[Dict[str, str]] = None) -> str:
    """
    Backward compatible function for existing code
    """
    try:
        # Convert simple lists to structured format for backward compatibility
        structured_participants = [
            {'name': p, 'mentions': 1, 'roles': []} if isinstance(p, str) else p 
            for p in participants
        ]
        structured_actions = [
            {'text': a, 'priority': 'unspecified', 'assignee': None} if isinstance(a, str) else a
            for a in action_items
        ]
        structured_dates = [
            {'text': d, 'context': '', 'normalized': d} if isinstance(d, str) else d
            for d in deadlines
        ]
        
        report_data = ReportData(
            participants=structured_participants,
            action_items=structured_actions,
            dates=structured_dates,
            summary=summary,
            transcript=original_transcript,
            title=title,
            metadata=metadata or {}
        )
        
        config = ReportConfig(export_format=ExportFormat.TXT)
        generator = get_report_generator()
        result = generator.generate_report(report_data, config)
        
        return result.get('content', 'Error generating report')
        
    except Exception as e:
        logger.exception("❌ Failed to generate backward compatible report")
        return f"Error generating meeting report: {str(e)}"

# Enhanced API functions
def generate_enhanced_report(participants: List[Dict],
                           action_items: List[Dict],
                           dates: List[Dict],
                           summary: str,
                           transcript: str,
                           speaker_segments: List[Tuple[str, str]] = None,
                           title: str = "Meeting Summary Report",
                           export_format: str = "txt",
                           include_visuals: bool = True,
                           include_statistics: bool = True) -> Dict:
    """
    Enhanced report generation with full feature set
    """
    try:
        format_enum = ExportFormat(export_format.lower())
    except ValueError:
        format_enum = ExportFormat.TXT
    
    report_data = ReportData(
        participants=participants,
        action_items=action_items,
        dates=dates,
        summary=summary,
        transcript=transcript,
        speaker_segments=speaker_segments or [],
        title=title
    )
    
    config = ReportConfig(
        export_format=format_enum,
        include_visuals=include_visuals,
        include_statistics=include_statistics,
        include_speaker_analysis=True,
        include_priority_analysis=True
    )
    
    generator = get_report_generator()
    return generator.generate_report(report_data, config)

def create_streamlit_report_ui(report_result: Dict):
    """
    Create Streamlit UI components for report display and download
    """
    if not report_result.get('success'):
        st.error(f"❌ Report generation failed: {report_result.get('error', 'Unknown error')}")
        return
    
    st.success("✅ Report generated successfully!")
    
    # Display analytics
    if report_result.get('analytics'):
        col1, col2 = st.columns(2)
        
        with col1:
            speaker_analysis = report_result['analytics'].get('speaker_analysis', {})
            if speaker_analysis.get('speaker_stats'):
                st.subheader("🎤 Speaker Analytics")
                st.metric("Total Speakers", speaker_analysis.get('total_speakers', 0))
                st.metric("Most Active", speaker_analysis.get('most_active', 'N/A'))
        
        with col2:
            priority_analysis = report_result['analytics'].get('priority_analysis', {})
            if priority_analysis:
                st.subheader("📊 Action Analytics") 
                st.metric("Total Actions", priority_analysis.get('total_actions', 0))
                st.metric("High Priority", priority_analysis.get('high_priority_count', 0))
    
    # Display charts
    charts = report_result.get('charts', {})
    if charts.get('speaker_contribution'):
        st.subheader("📈 Speaker Contribution")
        st.image(f"data:image/png;base64,{charts['speaker_contribution']}")
    
    if charts.get('priority_distribution'):
        st.subheader("📊 Priority Distribution")
        st.image(f"data:image/png;base64,{charts['priority_distribution']}")
    
    if charts.get('interactive_dashboard'):
        st.subheader("🔄 Interactive Dashboard")
        st.components.v1.html(charts['interactive_dashboard'], height=800)
    
    # Download button
    if report_result.get('content') and report_result.get('filename'):
        content = report_result['content']
        filename = report_result['filename']
        mime_type = report_result.get('mime_type', 'text/plain')
        
        # Convert to bytes if needed
        if isinstance(content, str):
            content = content.encode('utf-8')
        
        st.download_button(
            label=f"📥 Download {report_result['format'].upper()} Report",
            data=content,
            file_name=filename,
            mime=mime_type
        )
        
        # Preview for text formats
        if report_result['format'] in ['txt', 'json'] and isinstance(report_result['content'], (str, bytes)):
            preview_content = report_result['content']
            if isinstance(preview_content, bytes):
                preview_content = preview_content.decode('utf-8')
            
            with st.expander("👀 Preview Report"):
                if report_result['format'] == 'json':
                    st.json(json.loads(preview_content))
                else:
                    st.text_area("Report Preview", preview_content, height=400)

# Utility functions
def get_available_formats() -> List[str]:
    """Get list of available export formats"""
    return [fmt.value for fmt in ExportFormat]

def get_available_visual_styles() -> List[str]:
    """Get list of available visual styles"""
    return ['professional', 'modern', 'colorful']

def create_report_config(export_format: str = "txt",
                        include_visuals: bool = True,
                        visual_style: str = "professional",
                        include_transcript: bool = True,
                        **kwargs) -> ReportConfig:
    """Create report configuration with validation"""
    try:
        format_enum = ExportFormat(export_format.lower())
    except ValueError:
        format_enum = ExportFormat.TXT
        logger.warning(f"Invalid export format '{export_format}', using TXT")
    
    return ReportConfig(
        export_format=format_enum,
        include_visuals=include_visuals,
        visual_style=visual_style,
        include_transcript=include_transcript,
        **kwargs
    )


def generate_report_from_context(
    context: MeetingContext,
    export_format: str = "txt",
    include_visuals: bool = True
) -> Dict:
    """
    Pipeline adapter that converts MeetingContext into ReportData
    and generates a final report.
    """
    try:
        report_data = ReportData(
            participants=context.metadata.get("people", []),
            action_items=context.metadata.get("action_items", []),
            dates=context.metadata.get("dates", []),
            summary=context.metadata.get("summary", ""),
            transcript=context.raw_text,
            speaker_segments=context.speaker_segments,
            title=context.metadata.get("title", "Meeting Summary Report"),
            metadata=context.metadata
        )

        config = ReportConfig(
            export_format=ExportFormat(export_format.lower()),
            include_visuals=include_visuals,
            include_statistics=True,
            include_speaker_analysis=True,
            include_priority_analysis=True
        )

        generator = get_report_generator()
        return generator.generate_report(report_data, config)

    except Exception:
        logger.exception("❌ Failed to generate report from MeetingContext")
        return {
            "success": False,
            "error": "Report generation failed"
        }
 
